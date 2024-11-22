"""
written by Amir Nachum, from the legendary Machzor MEM ALEPH
"""
from abc import ABC, abstractmethod

from tkinter import Tk  # from tkinter import Tk for Python 3.x
from tkinter.filedialog import askdirectory

import os
import warnings

import numpy
import pandas as pd
import matplotlib.pyplot as plt
from tqdm import tqdm
import numpy as np

from docx.shared import Inches
from docx.oxml.shared import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from column_constants import SIGMAS

# wanted format of excel file
MASHOV_SAGZAB = ['name', 'personal example', 'involved', 'high standards', 'improving your training',
                 'accessible', 'respectful', 'familiarity', 'contributing conversations',
                 'contributing to personal development', 'desire to be under command',
                 'points to conserve', 'points to improve']
SOCIOMETRY = ["name", "knowing", "intrapersonal", "functioning in society", "leadership", "conduct", "academy",
              "applicative knowledge", "security", "responsibility", "excellence", "integrity", "daring", "mission",
              "courtesy", "points to conserve", "points to improve"]
VERBOSE = True

class PreProcess(ABC):
    def __init__(self):
        self.column_names = self.get_column_names()

    def run(self, input_dir_path, testing=False):
        self.input_dir = input_dir_path

        # sanity checks
        if not os.path.isdir(self.input_dir):
            raise Exception(f"the path is not a dir: {self.input_dir}")

        # process all excel files as dataframes
        dataframes = []
        for file_path in os.listdir(self.input_dir):
            # weird thing that it showed semi open files (with the $ symbol - denoting they are open)
            if "$" in file_path:
                continue

            # check if the file is an excel file
            if not file_path.endswith(".xlsx"):
                warnings.warn(f"non excel file in the dir: {file_path}")
                warnings.warn("skipping the file")
                continue

            df = pd.read_excel(self.input_dir + "\\" + file_path)
            df = self.remove_unwanted_data(df)
            df = df.rename(columns={df.columns[i]: self.column_names[i] for i in range(len(df.columns))})

            dataframes.append(df)

        # combine everything to one single big table
        combined_dataframe = self.combine_dataframes(dataframes)

        # drop all rows that the group by "name" in groupby is less than 3
        if not testing:
            combined_dataframe = combined_dataframe.groupby("name").filter(lambda x: len(x) >= 3)

        data_per_person = []
        for df in combined_dataframe.groupby(by="name"):
            data_per_person.append(df[1])
        return combined_dataframe, data_per_person

    def remove_unwanted_data(self, df: pd.DataFrame):
        # check if there is a timestamp column, and remove it
        first_column_name = df.columns[0]
        if first_column_name.startswith("חותמת"):
            df.drop(columns=first_column_name, inplace=True)

        # now check for segel column and remove it
        first_column_name = df.columns[0]
        if first_column_name.startswith("אני מסגל"):
            df.drop(columns=first_column_name, inplace=True)

        cols_to_remove = [col for col in df.columns if ("Unnamed" in col) or ("Column" in col)]
        if len(cols_to_remove) != 0:
            df.drop(columns=cols_to_remove, inplace=True)

        # remove empty lines
        df = df.dropna(axis=0, thresh=2)

        df = self.remove_platform_specific(df)
        return df

    @abstractmethod
    def remove_platform_specific(self, df: pd.DataFrame):
        pass

    @abstractmethod
    def get_column_names(self):
        pass

    @abstractmethod
    def combine_dataframes(self, dataframes) -> pd.DataFrame:
        pass


class Sagzab_PreProcess(PreProcess):
    def __init__(self):
        super().__init__()

    def get_column_names(self):
        column_names = MASHOV_SAGZAB
        return column_names

    def combine_dataframes(self, dataframes):
        ret_val = dataframes[0].copy()
        for df in dataframes[1:]:
            ret_val = ret_val.append(df, ignore_index=True)
        return ret_val

    def remove_platform_specific(self, df: pd.DataFrame):
        bad_columns = ["אני ממחלקת:"]
        if bad_columns[0] in df.columns:
            df.drop([bad_columns[0]], axis=1, inplace=True)
        return df


class Socio_PreProcess(PreProcess):
    def __init__(self):
        super().__init__()

    def get_column_names(self):
        return SOCIOMETRY

    def combine_dataframes(self, dataframes):
        ret_val = dataframes[0].copy()
        for df in dataframes[1:]:
            ret_val = ret_val.append(df, ignore_index=True)
        return ret_val

    def remove_platform_specific(self, df: pd.DataFrame):
        # remove avg row
        try:
            df.drop([0], inplace=True)
        except:
            pass
        return df


class Statistics(ABC):
    def __init__(self):
        pass

    def run(self, combined_data: pd.DataFrame):
        combined_data = combined_data.copy()
        # delete points to improve and points to conserve
        combined_data = combined_data.drop(columns=combined_data.columns[-2:])

        rows = []

        for df in combined_data.groupby("name"):
            person_name = df[0]
            df = df[1]
            columns = df.columns
            columns = columns.drop("name")
            for category in columns:
                col = df[category]
                # remove from col every non numeric value row
                col = col[col.apply(lambda x: isinstance(x, (int, np.int64, float, np.float64)))]
                mean = round(col.mean(), 2)
                std = round(col.std(), 2)
                row = [person_name, category, mean, std]
                rows.append(row)

        # what if I want to get more data, not open/close!!
        ret_val = pd.DataFrame(rows, columns=["name", "category", "mean", "std"])
        return ret_val


class Docx_helper(ABC):
    def __init__(self, file_format_path):
        self.file_format_path = file_format_path
        self.columns_names = self.get_columns_names()

    @abstractmethod
    def is_values(self, category: str) -> bool:
        pass

    @abstractmethod
    def get_columns_names(self):
        pass

    def sigma_text(self, sigma, is_values):
        # # mem bet values
        # values_big_threshold = 0.59
        # values_small_threshold = 0.38
        #
        # non_values_big_threshold = 1.2
        # non_values_small_threshold = 0.75

        # mem gimel values
        values_big_threshold = 0.57
        values_small_threshold = 0.39

        non_values_big_threshold = 1
        non_values_small_threshold = 0.65

        small, big = False, False
        if is_values:
            if sigma > values_big_threshold:
                big = True
            elif sigma < values_small_threshold:
                small = True
        else:
            if sigma > non_values_big_threshold:
                big = True
            elif sigma < non_values_small_threshold:
                small = True

        if big:
            ret_val = "sigma is top 15% (big)"
        elif small:
            ret_val = "sigma is lowest 15% (small)"
        else:
            ret_val = "sigma value is average"
        return ret_val

    def create_histogram(self, all_avgs, avg_total, avg_personal, std_personal, is_values=False, old_average=-1):
        fig = plt.figure()
        ax = plt.gca()

        if is_values:
            bins = [-0.125, 0.125, 0.375, 0.625, 0.875, 1.125, 1.375, 1.625, 1.875, 2.125, 2.375, 2.625, 2.875, 3.125]
            xticks = [0, 0.25, 0.5, 0.75, 1, 1.25, 1.5, 1.75, 2, 2.25, 2.5, 2.75, 3]
        else:
            bins = [-0.25, 0.25, 0.75, 1.25, 1.75, 2.25, 2.75, 3.25, 3.75, 4.25, 4.75, 5.25, 5.75, 6.25]
            xticks = [0, 0.5, 1, 1.5, 2, 2.5, 3, 3.5, 4, 4.5, 5, 5.5, 6]

        plt.xticks(xticks, fontsize=16)
        plt.hist(x=all_avgs, bins=bins, rwidth=0.9)
        plt.yticks(fontsize=16)

        # plot the average value of the specific person
        plt.axvline(avg_personal, color='red')

        # plot the average value of person form last year
        if old_average != -1:
            plt.axvline(old_average, color='green', linestyle="--")
            plt.text(0.01, 0.7, s="Red line - new result\nDashed line - last semester", fontsize=12, color='black',
                     transform=ax.transAxes)

        # plot the std of the specific person
        # the name column is the index, so we need the i'th column
        plt.hlines(y=sum(ax.get_ylim()) / 2, xmin=avg_personal - std_personal,
                   xmax=avg_personal + std_personal, color='red')
        # plt.text(0.01, 0.93, transform=ax.transAxes,
        #          s=r'$\sigma$' + f'={std_personal}\n{self.sigma_text(std_personal, is_values)}',
        #          fontsize=16, color='red')
        plt.text(0.01, 0.93, transform=ax.transAxes, s=r'$\sigma$' + f"={std_personal}", fontsize=16, color='red')
        plt.text(0.01, 0.89, s=self.sigma_text(std_personal, is_values), fontsize=16, color='red',
                 transform=ax.transAxes)
        SIGMAS.append(std_personal)

        plt.axvline(avg_total, color='black')
        secondary_ax = ax.secondary_xaxis("top")
        # plotting the value of the axvline on the histogram
        if abs(avg_personal - avg_total) < 0.2:
            diff = (0.2 - abs(avg_personal - avg_total)) / 2
            if avg_total > avg_personal:
                secondary_ax.set_xticks(ticks=[avg_personal - diff, avg_total + diff],
                                        labels=[f"{round(avg_personal, 2)}",
                                                f"{round(avg_total, 2)}"], rotation=60)
            else:
                secondary_ax.set_xticks(ticks=[avg_total - diff, avg_personal + diff],
                                        labels=[f"{round(avg_total, 2)}",
                                                f"{round(avg_personal, 2)}"], rotation=60)
        else:
            secondary_ax.set_xticks(ticks=[avg_personal, avg_total],
                                    labels=[f"{round(avg_personal, 2)}", f"{round(avg_total, 2)}"],
                                    rotation=60)

        for label in secondary_ax.get_xticklabels():
            label.set_fontsize(16)

        fig.set_size_inches(10, 5)
        plt.close(fig)
        # fig.show()
        # plt.show()
        return fig

    def create_word_file(self, hists, literal_data, person_name, n=None):
        doc = Document(self.file_format_path)

        title = doc.paragraphs[1]
        title.text += person_name
        if n is not None:
            title.text += f" (N={n})"
        title.runs[0] = "David"
        title.runs[0].underline = True

        # add the histograms to the table
        for i in range(len(hists)):
            # save the hist as png, so we can load to word as a picture
            cur_hist = hists[i]
            TMP_FILE_PATH = "tmp.png"
            cur_hist.savefig(TMP_FILE_PATH)

            # load png
            cell = doc.tables[0].cell(i, 1)
            cell.add_paragraph().add_run().add_picture(TMP_FILE_PATH, height=Inches(2))
            cell.paragraphs[1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # remove png as it is no longer needed
            os.remove(TMP_FILE_PATH)

        # points to conserve and point to improve
        conserve_cell = doc.tables[0].cell(len(hists), 1)
        improve_cell = doc.tables[0].cell(len(hists) + 1, 1)
        conserve_cell.add_paragraph("\n\n".join(literal_data[-2].dropna().array))
        improve_cell.add_paragraph("\n\n".join(literal_data[-1].dropna().array))

        # avoiding corrupting the word file
        # the id gets mixed with some of the things of the template file
        # for further reading - https://github.com/python-openxml/python-docx/issues/455
        # or need version better than 0.8.7
        try:
            docPrs = doc._part._element.findall('.//' + qn('wp:docPr'))
            for docPr in docPrs:
                docPr.set('id', str(int(docPr.get('id')) + 100000))
        except:
            pass

        return doc

    def run(self, combined_df: pd.DataFrame, stats_df: pd.DataFrame, old_stats_df=None):
        # create word file for every person
        ret_val = []
        for df in tqdm(combined_df.groupby("name"), disable=not VERBOSE):
            person_name = df[0]
            df = df[1]
            conserve_points, improve_points = df[df.columns[-2]], df[df.columns[-1]]

            # TODO - voodoo code to get only numerical columns
            numerical_columns = df.columns.drop("name")[:-2]
            stats_per_person = stats_df[stats_df["name"] == person_name]
            hists = []
            for category in numerical_columns:
                # calculate old value
                old_average = -1
                if old_stats_df is not None:
                    old_res = \
                        old_stats_df[(old_stats_df["name"] == person_name) & (old_stats_df["category"] == category)][
                            "mean"]
                    if len(old_res) > 0:
                        old_average = old_res.array[0]

                combined_col = combined_df[category]
                # remove from col every non numeric value row
                combined_col = combined_col[combined_col.apply(lambda x: isinstance(x, (int, np.int64, float, np.float64)))]

                avg_total = combined_col.mean()
                avg_personal = stats_per_person[stats_per_person["category"] == category]["mean"].values[0]
                std_personal = stats_per_person[stats_per_person["category"] == category]["std"].values[0]
                all_avgs = stats_df[stats_df["category"] == category]["mean"]

                hist = self.create_histogram(all_avgs, avg_total, avg_personal, std_personal, self.is_values(category),
                                             old_average=old_average)
                hists.append(hist)

            N = df.shape[0]
            word_file = self.create_word_file(hists, [conserve_points, improve_points], person_name, n=N)
            title_to_save = f"{person_name} (N={N})"
            ret_val.append([word_file, title_to_save])

        return ret_val


class Docx_sagzab(Docx_helper):
    def __init__(self, file_format_path):
        super().__init__(file_format_path)

    def is_values(self, category: str) -> bool:
        return False

    def get_columns_names(self):
        column_names = MASHOV_SAGZAB
        return column_names


class Docx_Socio(Docx_helper):
    def __init__(self, file_format_path):
        super().__init__(file_format_path)

    def is_values(self, category: str) -> bool:
        if category in SOCIOMETRY[-8:-2]:
            return True
        return False

    def get_columns_names(self):
        column_names = SOCIOMETRY
        return column_names


class RunObj():
    def __init__(self):
        pass

    @abstractmethod
    def get_preprocess_obj(self):
        pass

    @abstractmethod
    def get_docx_obj(self):
        pass

    @abstractmethod
    def get_file_format_path(self):
        pass

    def run(self, output_dir, old_stats_path=None, testing=False):
        if not os.path.isdir(output_dir):
            os.mkdir(output_dir)

        # ask user to choose the input directory (where all of the excels are)
        Tk().withdraw()
        input_dir_path = askdirectory()  # close files window before using

        preprocess = self.get_preprocess_obj()()
        combined_df, data_per_person_list = preprocess.run(input_dir_path, testing=testing)

        # save raw_data_per_person and combined data

        combined_df.to_excel(output_dir + "\\combined_data.xlsx", index=False)
        raw_data_dir_path = output_dir + "\\raw_data"
        if not os.path.isdir(raw_data_dir_path):
            os.mkdir(raw_data_dir_path)
        for df in data_per_person_list:
            N = df.shape[0]
            person_name = df["name"].iloc[0]
            file_title = raw_data_dir_path + f"\\{person_name} (N={N}).xlsx"
            file_title = file_title.replace('"', '')
            file_title = file_title.replace("'", '')
            df.to_excel(file_title, index=False)

        import numpy as np
        # create stats excel, and save
        stat_obj = Statistics()
        stats_df = stat_obj.run(combined_df)

        stats_df["n_sigma"] = stats_df.groupby("category")["mean"].apply(lambda x: (x - numpy.mean(x)) / numpy.std(x))
        stats_df.to_excel(output_dir + "\\stats_excel.xlsx", index=False)

        # get the old data
        old_stats_df = None
        if old_stats_path is not None:
            old_stats_df = pd.read_excel(old_stats_path, header=0, engine="openpyxl")

        # create word files
        docx_helper = self.get_docx_obj()(self.get_file_format_path())
        ret_val = docx_helper.run(combined_df, stats_df, old_stats_df=old_stats_df)
        import numpy as np
        np.savetxt("memdalet_sigmas.csv", SIGMAS, delimiter=',')

        # save word files
        print("making word files")
        word_output_dir = output_dir + "\\word"
        if not os.path.isdir(word_output_dir):
            os.mkdir(word_output_dir)
        for word_file, title in ret_val:
            title = title.replace('"', '')
            title = title.replace("'", '')
            word_file.save(word_output_dir + f"\\{title}.docx")


class RunSocio(RunObj):
    def __init__(self):
        super().__init__()
        self.file_format_path = "socio_format.docx"

    def get_docx_obj(self):
        return Docx_Socio

    def get_preprocess_obj(self):
        return Socio_PreProcess

    def get_file_format_path(self):
        return self.file_format_path


class RunSagzab(RunObj):
    def __init__(self):
        super().__init__()
        self.file_format_path = "mashov_mefakdim_format.docx"

    def get_docx_obj(self):
        return Docx_sagzab

    def get_preprocess_obj(self):
        return Sagzab_PreProcess

    def get_file_format_path(self):
        return self.file_format_path


if __name__ == '__main__':
    VERBOSE = True
    run_socio = True
    output_dir_socio = r"sociometry_output"
    output_dir_sagab_sagaz = r"sagaz_output"
    old_stats_path = None  # can be None
    if run_socio:
        run = RunSocio()
    else:
        run = RunSagzab()

    if run_socio:
        run.run(output_dir_socio, old_stats_path=old_stats_path, testing=True)
    else:
        run.run(output_dir_sagab_sagaz, testing=True)
