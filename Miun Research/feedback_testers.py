import pandas as pd
import math
import plotly.express as px
import os
from tqdm import tqdm
from docx.shared import Inches
from docx.oxml.shared import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from sklearn.metrics import roc_curve, auc
from scipy import stats
import dataframe_image as dfi

# for not making warnings
pd.options.mode.chained_assignment = None  # default='warn'

FORMAT_WORD_PATH = r"Formats\feedback_format.docx"
BASE_SAVE_PATH = r"Miun Research\Results"


def get_df_of_tests(path_to_feedbacks=r"C:\Users\t9028387\PycharmProjects\TalpiotCodes\Miun Research\משובי מעריכים00 .xlsx"):
    """
    get the data frame of the tests of last year
    :return:
    """
    df = pd.read_excel(path_to_feedbacks, header=0, engine="openpyxl")

    for index, row in df.iterrows():
        try:
            if math.isnan(df.at[index, "ממוצע אקדמי"]):
                df.at[index, "ממוצע אקדמי"] = None
            else:
                val = float(df.at[index, "ממוצע אקדמי"])
                df.at[index, "ממוצע אקדמי"] = val
        except:
            df.at[index, "ממוצע אקדמי"] = None

    df = df.replace("גיורא שורץ", "גיורא שוורץ")
    df = df.replace("דוד פוטריאן", "דוד פוטוריאן")
    df = df.replace("נועה  קורן", "נועה קורן")
    df = df.replace("עדו טרבצקי", "עדו טריביצקי")
    df = df.replace("איתי הולצמן ", "איתי הולצמן")
    df = df.replace("עידו ברקוביץ", "עידן ברקוביץ")
    df = df.replace("אריאל פרסקו", "אריאל פרסיקו")
    df = df.replace("עודד", "עודד ניר")
    df = df.replace("אחי אור", "אחי אור וינגרטן")
    df = df.replace('חק"ב', "חקב")
    df["passed"] = df["התקבל/לא התקבל"] == "התקבל"

    for tester in df["מעריך"].unique():
        tester_path = fr"C:\Users\t9028387\PycharmProjects\TalpiotCodes\Miun Research\Results\{tester}"
        if not os.path.exists(tester_path):
            os.mkdir(tester_path)

    columns_dict = {'מעריך': "tester",
                    'סוג מבחן': "kind",
                    'מסכם': "miun grade",
                    'התקבל/לא התקבל': "miun status",
                    'סטטוס בהכשרה (נשאר/עלה ועדה/פרש/אי התאמת המערכת)': 'status',
                    'ממוצע אקדמי': "academy average",
                    "תז מוערך": "id",
                    "סדרה": "sidra number",
                    "קבוצה": "group",
                    "מוערך": "malshab_first_name",
                    "שם משפחה מוערך": "malshab_last_name",
                    }

    df = df.rename(columns=columns_dict)

    df["is in talpiot"] = (df["status"] == "בהכשרה") | (df["status"] == "בדחש")

    # add pass/not passed
    have_grades_df = df[df["academy average"].notnull()]

    AUC_dict_by_tester = {}
    for tester in df["tester"].unique():
        tester_df = df[(df["tester"] == tester)]
        fpr, tpr, thresholds = roc_curve(tester_df["is in talpiot"], tester_df["miun grade"])
        AUC_dict_by_tester[tester] = auc(fpr, tpr)

    AUC_df = pd.DataFrame.from_dict(AUC_dict_by_tester, orient='index', columns=["AUC"])

    return {"df": df, "have_grades_df": have_grades_df, "AUC_df": AUC_df}


def create_word_feedback(dict_df, format_path=FORMAT_WORD_PATH, base_save_path=BASE_SAVE_PATH):
    # create word files
    for tester_name in tqdm(dict_df["df"]["tester"].unique()):
        doc = create_word_file(format_path, dict_df, tester_name)
        path = os.path.join(base_save_path, tester_name, f"{tester_name}_feedback.docx")
        doc.save(path)


def create_word_file(format_path, dict_df, tester_name):
    doc = Document(format_path)

    title = doc.paragraphs[1]
    title.text += tester_name
    title.runs[0] = "David"
    title.runs[0].underline = True

    # for every cell calls the function that returns the graphs and texts
    params_dict = {"df": dict_df["df"], "tester_name": tester_name, "have_grades_df": dict_df["have_grades_df"],
                   "AUC_df": dict_df["AUC_df"]}
    for cell_num in WORD_CELLS_TO_FUNCTIONS.keys():
        feedback_func = WORD_CELLS_TO_FUNCTIONS[cell_num]

        # in case there is a graph and texts to add
        if feedback_func is not None:
            ret_dict = feedback_func(params_dict)
            graphs = ret_dict["graphs"]
            texts = ret_dict["texts"]
            tables = ret_dict["tables"]

            num_paragraphs = 0
            cell = doc.tables[0].cell(cell_num, 1)
            # add the graphs
            for graph in graphs:
                # save the plotly graph temporary
                tmp_file_path = "tmp.png"
                graph.write_image(tmp_file_path)

                # load png
                cell.add_paragraph().add_run().add_picture(tmp_file_path, height=Inches(4))
                num_paragraphs += 1
                cell.paragraphs[num_paragraphs].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # remove png as it is no longer needed
                os.remove(tmp_file_path)

            # add the texts
            for text in texts:
                cell.add_paragraph(text)
                num_paragraphs += 1

            # add the tables
            for table in tables:
                # save the plotly graph temporary
                tmp_file_path = "tmp.png"
                dfi.export(table, tmp_file_path)

                # load png
                cell.add_paragraph().add_run().add_picture(tmp_file_path, width=Inches(5))
                num_paragraphs += 1
                cell.paragraphs[num_paragraphs].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # remove png as it is no longer needed
                os.remove(tmp_file_path)

    # avoiding corrupting the word file
    # the id gets mixed with some of the things of the template file
    # for further reading - https://github.com/python-openxml/python-docx/issues/455
    docPrs = doc._part._element.findall('.//' + qn('wp:docPr'))
    for docPr in docPrs:
        docPr.set('id', str(int(docPr.get('id')) + 100000))

    return doc


def toughness_graphs(params_dict):
    """
    create graphs of toughness of tester
    :param params_dict:
    :return: graphs, texts
    """
    df = params_dict["df"]
    tester_name = params_dict["tester_name"]

    figs = []
    texts = []

    try:
        passed_toughness = df[df["passed"] == True].groupby("tester")["miun grade"].mean()
        fig = px.histogram(passed_toughness,
                           title="Average miun grade of tester histogram - for malshabs that passed<br><br>(the line is your average)")

        fig.add_vline(x=passed_toughness.at[tester_name])

        figs.append(fig)
    except:
        print(f"failed toughness_graphs {tester_name} - passed malshabs")

    try:
        failed_toughness = df[df["passed"] == False].groupby("tester")["miun grade"].mean()
        fig = px.histogram(failed_toughness,
                           title="Average miun grade of tester histogram - for malshabs that failed<br><br>(the line is your average)")

        fig.add_vline(x=failed_toughness.at[tester_name])

        figs.append(fig)
    except:
        print(f"failed toughness_graphs {tester_name} - failed malshabs")

    ret_dict = {"graphs": figs, "texts": texts, "tables": []}
    return ret_dict


def academy_by_miun_grades(params_dict):
    """
    create graphs of toughness of tester
    :param params_dict:
    :return: graphs, texts
    """
    df = params_dict["df"]
    have_grades_df = params_dict["have_grades_df"]
    tester_name = params_dict["tester_name"]

    figs = []
    texts = []

    custom_grades_df = have_grades_df
    custom_grades_df["is_my_examinee"] = (custom_grades_df["tester"] == tester_name)
    fig = px.scatter(custom_grades_df[custom_grades_df["is_my_examinee"] == True], x="miun grade", y="academy average",
                     symbol="kind", color="is_my_examinee", symbol_sequence=['circle-open', 'x'],
                     title="academy average by miun grade of my examinees")
    fig.update_xaxes(range=[0, 6])

    figs.append(fig)

    ret_dict = {"graphs": figs, "texts": texts, "tables": []}
    return ret_dict


def get_table_of_tester(params_dict):
    """
    :param params_dict:
    :return:
    """
    df = params_dict["df"]
    have_grades_df = params_dict["have_grades_df"]
    tester_name = params_dict["tester_name"]

    tables = []

    relevant_columns = ['tester', 'sidra number', 'group', 'id', 'malshab_last_name',
                        'malshab_first_name', 'kind', 'miun grade', 'miun status',
                        'status', 'academy average', "passed"]
    tester_df = df[relevant_columns][df["tester"] == tester_name]
    tables.append(tester_df)

    ret_dict = {"graphs": [], "texts": [], "tables": tables}
    return ret_dict


def get_roc_curve(params_dict):
    """
    return text of 2 values of success
    :param params_dict:
    :return:
    """
    df = params_dict["df"]
    have_grades_df = params_dict["have_grades_df"]
    tester_name = params_dict["tester_name"]
    AUC_df = params_dict["AUC_df"]

    # ROC curve

    tester_df = df[(df["tester"] == tester_name)]
    fpr, tpr, thresholds = roc_curve(tester_df["is in talpiot"], tester_df["miun grade"])

    fig = px.area(
        x=fpr, y=tpr,
        title=f'ROC Curve (AUC={auc(fpr, tpr):.4f})',
        labels=dict(x='False Positive Rate', y='True Positive Rate'),
        width=700, height=500
    )
    fig.add_shape(
        type='line', line=dict(dash='dash'),
        x0=0, x1=1, y0=0, y1=1
    )

    fig.update_yaxes(scaleanchor="x", scaleratio=1)
    fig.update_xaxes(constrain='domain')

    # AUC histogram
    fig2 = px.histogram(AUC_df, title="AUC histogram<br><br>(the line is your AUC)")
    fig2.add_vline(x=AUC_df["AUC"].at[tester_name])

    text = "The ROC curve and the AUC is generated from your grades in the miun compared to the real data of 'is in Talpiot'\n"
    text += "Note: if malshab passed the tests but kicked out/leaves Talpiot - he is considered as not in Talpiot\n"
    text += "Note: take carefully the result and the AUC - it's not based on big data\n"
    text += "Note: if none of your malshabs is passed then the ROC is meaningless"
    ret_dict = {"graphs": [fig, fig2], "texts": [text], "tables": []}
    return ret_dict


WORD_CELLS_TO_FUNCTIONS = {
    0: None,
    1: get_table_of_tester,
    2: toughness_graphs,
    3: academy_by_miun_grades,
    4: get_roc_curve
}

if __name__ == "__main__":

    get_df_of_miun_results = True
    create_feedback_files = True

    if get_df_of_miun_results:
        dict_df = get_df_of_tests()

    if create_feedback_files:
        create_word_feedback(dict_df)
