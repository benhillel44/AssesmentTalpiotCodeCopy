from abc import ABC, abstractmethod
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tqdm import tqdm
import os
from docx.oxml.ns import qn
from column_constants import SIGMAS
from docxtpl import DocxTemplate

class Docx_helper(ABC):
    def __init__(self, file_format_path, word_output_dir):
        self.file_format_path = file_format_path
        self.word_output_dir = word_output_dir


    @abstractmethod
    def is_values(self, category: str) -> bool:
        pass


    @abstractmethod
    def get_columns_names(self):
        pass


    def sigma_text(self, sigma, is_values):
        # # mem bet values
        # values_big_threshold = 0.59
        # values_small_threshold = 0.38
        #
        # non_values_big_threshold = 1.2
        # non_values_small_threshold = 0.75

        # mem gimel values
        values_big_threshold = 0.57
        values_small_threshold = 0.39

        non_values_big_threshold = 1
        non_values_small_threshold = 0.65

        small, big = False, False
        if is_values:
            if sigma > values_big_threshold:
                big = True
            elif sigma < values_small_threshold:
                small = True
        else:
            if sigma > non_values_big_threshold:
                big = True
            elif sigma < non_values_small_threshold:
                small = True

        if big:
            ret_val = "sigma is top 15% (big)"
        elif small:
            ret_val = "sigma is lowest 15% (small)"
        else:
            ret_val = "sigma value is average"
        return ret_val


    def create_histogram(self, all_avgs, avg_total, avg_personal, std_personal, is_values=False, old_average=-1):
        fig = plt.figure()
        ax = plt.gca()

        if is_values:
            bins = np.arange(-0.125, 3.375, 0.25)
            xticks = np.arange(0, 3.25, 0.25)
        else:
            bins = np.arange(-0.25, 6.75, 0.5)
            xticks = np.arange(0, 6.5, 0.5)

        plt.xticks(xticks, fontsize=16)
        plt.hist(x=all_avgs, bins=bins, rwidth=0.9)
        plt.yticks(fontsize=16)

        # plot the average value of the specific person
        plt.axvline(avg_personal, color='red')

        # plot the average value of person form last year
        if old_average != -1:
            plt.axvline(old_average, color='green', linestyle="--")
            plt.text(0.01, 0.7, s="Red line - new result\nDashed line - last semester", fontsize=12, color='black',
                     transform=ax.transAxes)

        # plot the std of the specific person
        # the name column is the index, so we need the i'th column
        plt.hlines(y=sum(ax.get_ylim()) / 2, xmin=avg_personal - std_personal,
                   xmax=avg_personal + std_personal, color='red')
        # plt.text(0.01, 0.93, transform=ax.transAxes,
        #          s=r'$\sigma$' + f'={std_personal}\n{self.sigma_text(std_personal, is_values)}',
        #          fontsize=16, color='red')
        plt.text(0.01, 0.93, transform=ax.transAxes, s=r'$\sigma$' + f"={std_personal}", fontsize=16, color='red')
        plt.text(0.01, 0.89, s=self.sigma_text(std_personal, is_values), fontsize=16, color='red',
                 transform=ax.transAxes)

        plt.axvline(avg_total, color='black')
        secondary_ax = ax.secondary_xaxis("top")
        # plotting the value of the axvline on the histogram
        if abs(avg_personal - avg_total) < 0.2:
            diff = (0.2 - abs(avg_personal - avg_total)) / 2
            if avg_total > avg_personal:
                secondary_ax.set_xticks(ticks=[avg_personal - diff, avg_total + diff],
                                        labels=[f"{round(avg_personal, 2)}",
                                                f"{round(avg_total, 2)}"], rotation=60)
            else:
                secondary_ax.set_xticks(ticks=[avg_total - diff, avg_personal + diff],
                                        labels=[f"{round(avg_total, 2)}",
                                                f"{round(avg_personal, 2)}"], rotation=60)
        else:
            secondary_ax.set_xticks(ticks=[avg_personal, avg_total],
                                    labels=[f"{round(avg_personal, 2)}", f"{round(avg_total, 2)}"],
                                    rotation=60)

        for label in secondary_ax.get_xticklabels():
            label.set_fontsize(16)

        fig.set_size_inches(10, 5)
        plt.close(fig)
        # fig.show()
        # plt.show()
        return fig
    

    def insert_classifications(self, classification_df, format_file_name):
        conserve_names = [("Interpersonal Skills", "יכולות בין-אישיות"),\
                       ("Intrapersonal Skills", "יכולות תוך-אישיות"),\
                       ("Professionalism", "מקצועיות"),\
                       ("conduct", "התנהלות"),\
                       ("Leadership", "מנהיגות"),\
                       ("Other", "אחר")]

        improve_names = [("Interpersonal Skills2", "יכולות בין-אישיות"),\
                       ("Intrapersonal Skills2", "יכולות תוך-אישיות"),\
                       ("Professionalism2", "מקצועיות"),\
                       ("conduct2", "התנהלות"),\
                       ("Leadership2", "מנהיגות"),\
                       ("Other2", "אחר")]
        
        doc = DocxTemplate(template_file=format_file_name)
        context = {'conserve_classifications':[], 'improve_classifications':[]}

        for column, word_name in conserve_names:
            list_of_sentences = []
            if column not in classification_df:
                continue

            for i in range(len(classification_df[column])):
                if classification_df[column][i] in ["True", True, "TRUE", "true"]:
                    list_of_sentences.append({'name':classification_df['Original_conserve'][i]})

            if len(list_of_sentences) > 0:
                class_dict = {'name':f"(N={len(list_of_sentences)}) "+word_name}
                class_dict["bullets"] = list_of_sentences
                
                context["conserve_classifications"].append(class_dict)
        
        for column, word_name in improve_names:
            list_of_sentences = []
            if column not in classification_df:
                continue

            for i in range(len(classification_df[column])):
                if classification_df[column][i] in ["True", True, "TRUE", "true"]:
                    list_of_sentences.append({'name':classification_df['Original_improve'][i]})

            if len(list_of_sentences) > 0:
                class_dict = {'name':f"(N={len(list_of_sentences)}) "+word_name}
                class_dict["bullets"] = list_of_sentences
                
                context["improve_classifications"].append(class_dict)
        
        doc.render(context=context)
        doc.save(format_file_name)


    def create_word_file(self, hists, classification_df, person_name, n=None):
        title_to_save = f"{person_name} (N={n})".replace('"', '').replace("'", '')+".docx"
        path_to_save = os.path.join(self.word_output_dir,title_to_save)

        doc = Document(self.file_format_path)

        title = doc.paragraphs[1]
        title.text += person_name
        if n is not None:
            title.text += f" (N={n})"
        title.runs[0] = "David"
        title.runs[0].underline = True

        # add the histograms to the table
        for i in range(len(hists)):
            # save the hist as png, so we can load to word as a picture
            cur_hist = hists[i]
            TMP_FILE_PATH = "tmp.png"
            cur_hist.savefig(TMP_FILE_PATH)

            # load png
            cell = doc.tables[0].cell(i, 1)
            cell.add_paragraph().add_run().add_picture(TMP_FILE_PATH, height=Inches(2))
            cell.paragraphs[1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # remove png as it is no longer needed
            os.remove(TMP_FILE_PATH)

        # avoiding corrupting the word file
        # the id gets mixed with some of the things of the template file
        # for further reading - https://github.com/python-openxml/python-docx/issues/455
        # or need version better than 0.8.7
        try:
            docPrs = doc._part._element.findall('.//' + qn('wp:docPr'))
            for docPr in docPrs:
                docPr.set('id', str(int(docPr.get('id')) + 100000))
        except:
            pass

        doc.save(path_to_save)

        if classification_df is not None:
            self.insert_classifications(classification_df, path_to_save)


    def run_word_creation(self,
            combined_df: pd.DataFrame,
            stats_df: pd.DataFrame,
            name_to_classification:dict=None,
            old_stats_df=None,
            verbose: bool = True,
            start_cadet: str = None):

        reached_start_cadet = start_cadet is None
        # create word file for every person
        for df in tqdm(combined_df.groupby("name"), disable=not verbose):
            person_name = df[0]
            if not reached_start_cadet:
                if person_name == start_cadet:
                    reached_start_cadet = True
                else:
                    continue
            
            df = df[1]

            # TODO - voodoo code to get only numerical columns
            numerical_columns = df.columns.drop("name")[:-2]
            stats_per_person = stats_df[stats_df["name"] == person_name]
            hists = []
            for category in numerical_columns:
                # calculate old value
                old_average = -1
                if old_stats_df is not None:
                    # does old_stats_df contain the person?
                    if len(old_stats_df[old_stats_df["name"] == person_name]) == 0:
                        print(f"Person {person_name} not found in old stats!\n"
                              f"probably someone changed their name in the raw data excel file\n")

                    old_res = \
                        old_stats_df[(old_stats_df["name"] == person_name) & (old_stats_df["category"] == category)][
                            "mean"]
                    if len(old_res) > 0:
                        old_average = old_res.array[0]

                combined_col = combined_df[category]
                # remove from col every non numeric value row
                combined_col = combined_col[combined_col.apply(lambda x: isinstance(x, (int, np.int64, float, np.float64)))]

                avg_total = combined_col.mean()
                avg_personal = stats_per_person[stats_per_person["category"] == category]["mean"].values[0]
                std_personal = stats_per_person[stats_per_person["category"] == category]["std"].values[0]
                SIGMAS.append(std_personal)
                all_avgs = stats_df[stats_df["category"] == category]["mean"]

                hist = self.create_histogram(all_avgs, avg_total, avg_personal, std_personal, self.is_values(category),
                                             old_average=old_average)
                hists.append(hist)

            N = df.shape[0]
            print(f"Creating word file for {person_name} (N={N})")
            df = name_to_classification[person_name] if name_to_classification is not None else None
            self.create_word_file(hists, df, person_name, n=N)
